"""
Log Sentinel v2.0 - Report Exporter
====================================
Exportação de relatórios em múltiplos formatos.

Author: Duarte Cunha (Nº 2024271)
ISTEC - Instituto Superior de Tecnologias Avançadas de Lisboa
Ano Letivo: 2025/2026

Formatos suportados:
- PDF (relatório profissional)
- CSV (dados tabulares)
- JSON (dados estruturados)
- DOCX (documento Word)
"""

import json
import csv
from datetime import datetime
from pathlib import Path
from typing import List, Dict, Any, Optional
from dataclasses import asdict

# PDF
try:
    from reportlab.lib import colors
    from reportlab.lib.pagesizes import A4
    from reportlab.lib.styles import getSampleStyleSheet, ParagraphStyle
    from reportlab.lib.units import cm, mm
    from reportlab.platypus import (
        SimpleDocTemplate, Paragraph, Spacer, Table, TableStyle,
        PageBreak, Image, HRFlowable
    )
    from reportlab.lib.enums import TA_CENTER, TA_LEFT, TA_RIGHT
    REPORTLAB_AVAILABLE = True
except ImportError:
    REPORTLAB_AVAILABLE = False

# DOCX
try:
    from docx import Document
    from docx.shared import Inches, Pt, Cm
    from docx.enum.text import WD_ALIGN_PARAGRAPH
    from docx.enum.table import WD_TABLE_ALIGNMENT
    DOCX_AVAILABLE = True
except ImportError:
    DOCX_AVAILABLE = False


class ReportExporter:
    """
    Exportador de relatórios multi-formato.
    
    Gera relatórios profissionais com:
    - Resumo executivo
    - Estatísticas e métricas
    - Lista de anomalias
    - Gráficos e visualizações
    - Recomendações
    """
    
    def __init__(self):
        self.author = "Duarte Cunha"
        self.student_id = "2024271"
        self.institution = "ISTEC - Instituto Superior de Tecnologias Avançadas de Lisboa"
        self.app_name = "Log Sentinel"
        self.version = "2.0"
    
    def get_supported_formats(self) -> List[str]:
        """Retorna formatos suportados."""
        formats = ['json', 'csv']
        if REPORTLAB_AVAILABLE:
            formats.append('pdf')
        if DOCX_AVAILABLE:
            formats.append('docx')
        return formats
    
    def _convert_anomalies(self, anomalies: List) -> List[Dict]:
        """Converte lista de anomalias para lista de dicts."""
        result = []
        for a in anomalies:
            if isinstance(a, dict):
                result.append(a)
            elif hasattr(a, 'to_dict'):
                result.append(a.to_dict())
            else:
                # Conversão manual
                d = {}
                
                # Atributos comuns
                if hasattr(a, 'anomaly_type'):
                    val = a.anomaly_type
                    d['type'] = val.value if hasattr(val, 'value') else str(val)
                
                if hasattr(a, 'severity'):
                    val = a.severity
                    d['severity'] = val.value if hasattr(val, 'value') else str(val)
                
                if hasattr(a, 'source_ip'):
                    d['source_ip'] = a.source_ip
                
                if hasattr(a, 'target'):
                    d['target'] = a.target
                
                if hasattr(a, 'detail'):
                    d['detail'] = a.detail
                
                if hasattr(a, 'timestamp'):
                    val = a.timestamp
                    d['timestamp'] = val.isoformat() if hasattr(val, 'isoformat') else str(val) if val else None
                
                if hasattr(a, 'evidence'):
                    d['evidence'] = a.evidence
                
                if hasattr(a, 'log_file'):
                    d['log_file'] = a.log_file
                
                if hasattr(a, 'ml_score'):
                    d['score'] = a.ml_score
                elif hasattr(a, 'score'):
                    d['score'] = a.score
                
                result.append(d)
        return result
    
    def export_json(self, anomalies: List, filepath: str, stats: Dict = None) -> bool:
        """Exporta para JSON."""
        anomalies = self._convert_anomalies(anomalies)
        return self._export_json(anomalies, stats or {}, Path(filepath))
    
    def export_csv(self, anomalies: List, filepath: str, stats: Dict = None) -> bool:
        """Exporta para CSV."""
        anomalies = self._convert_anomalies(anomalies)
        return self._export_csv(anomalies, stats or {}, Path(filepath))
    
    def export_pdf(self, anomalies: List, stats: Dict, filepath: str, log_file: str = None) -> bool:
        """Exporta para PDF."""
        if not REPORTLAB_AVAILABLE:
            return False
        anomalies = self._convert_anomalies(anomalies)
        return self._export_pdf(anomalies, stats, Path(filepath), f"Relatório - {log_file or 'Análise'}")
    
    def export_docx(self, anomalies: List, stats: Dict, filepath: str, log_file: str = None) -> bool:
        """Exporta para DOCX."""
        if not DOCX_AVAILABLE:
            return False
        anomalies = self._convert_anomalies(anomalies)
        return self._export_docx(anomalies, stats, Path(filepath), f"Relatório - {log_file or 'Análise'}")
    
    def export(self, anomalies: List[Dict], stats: Dict, 
               filepath: str, format: str = "pdf",
               title: str = "Relatório de Análise de Segurança") -> bool:
        """
        Exporta relatório no formato especificado.
        
        Args:
            anomalies: Lista de anomalias detetadas
            stats: Estatísticas da análise
            filepath: Caminho do ficheiro de saída
            format: Formato (pdf, csv, json, docx)
            title: Título do relatório
            
        Returns:
            True se exportado com sucesso
        """
        filepath = Path(filepath)
        filepath.parent.mkdir(parents=True, exist_ok=True)
        
        format = format.lower()
        
        if format == "pdf":
            return self._export_pdf(anomalies, stats, filepath, title)
        elif format == "csv":
            return self._export_csv(anomalies, stats, filepath)
        elif format == "json":
            return self._export_json(anomalies, stats, filepath)
        elif format == "docx":
            return self._export_docx(anomalies, stats, filepath, title)
        else:
            raise ValueError(f"Formato não suportado: {format}")
    
    # === PDF Export ===
    
    def _export_pdf(self, anomalies: List[Dict], stats: Dict, 
                    filepath: Path, title: str) -> bool:
        """Exporta relatório em PDF."""
        if not REPORTLAB_AVAILABLE:
            raise ImportError("reportlab não disponível")
        
        doc = SimpleDocTemplate(
            str(filepath),
            pagesize=A4,
            rightMargin=2*cm,
            leftMargin=2*cm,
            topMargin=2*cm,
            bottomMargin=2*cm
        )
        
        styles = getSampleStyleSheet()
        
        # Custom styles
        styles.add(ParagraphStyle(
            name='CustomTitle',
            parent=styles['Heading1'],
            fontSize=24,
            spaceAfter=30,
            textColor=colors.HexColor('#1e3a5f'),
            alignment=TA_CENTER
        ))
        
        styles.add(ParagraphStyle(
            name='CustomHeading',
            parent=styles['Heading2'],
            fontSize=14,
            spaceBefore=20,
            spaceAfter=10,
            textColor=colors.HexColor('#2563eb'),
        ))
        
        styles.add(ParagraphStyle(
            name='CustomBody',
            parent=styles['Normal'],
            fontSize=10,
            spaceAfter=8,
        ))
        
        # Build content
        story = []
        
        # Header
        story.append(Paragraph("🦉 LOG SENTINEL", styles['CustomTitle']))
        story.append(Paragraph(title, styles['Heading2']))
        story.append(Spacer(1, 10))
        
        # Meta info
        meta_data = [
            ['Data do Relatório:', datetime.now().strftime('%d/%m/%Y %H:%M')],
            ['Autor:', f"{self.author} (Nº {self.student_id})"],
            ['Instituição:', self.institution],
            ['Versão:', f"{self.app_name} v{self.version}"],
        ]
        meta_table = Table(meta_data, colWidths=[4*cm, 12*cm])
        meta_table.setStyle(TableStyle([
            ('FONTSIZE', (0, 0), (-1, -1), 9),
            ('TEXTCOLOR', (0, 0), (0, -1), colors.gray),
            ('ALIGN', (0, 0), (0, -1), 'RIGHT'),
            ('ALIGN', (1, 0), (1, -1), 'LEFT'),
        ]))
        story.append(meta_table)
        story.append(Spacer(1, 20))
        
        # Linha separadora
        story.append(HRFlowable(width="100%", thickness=1, color=colors.HexColor('#3b82f6')))
        story.append(Spacer(1, 20))
        
        # Resumo Executivo
        story.append(Paragraph("1. Resumo Executivo", styles['CustomHeading']))
        
        total = stats.get('total_anomalies', stats.get('anomalies_detected', 0))
        critical = stats.get('by_severity', {}).get('CRITICAL', 0)
        high = stats.get('by_severity', {}).get('HIGH', 0)
        
        summary = f"""
        Esta análise processou <b>{stats.get('entries_processed', 'N/A')}</b> entradas de log
        e identificou <b>{total}</b> anomalias de segurança.
        <br/><br/>
        <b>Severidade das ameaças:</b><br/>
        • Críticas: {critical}<br/>
        • Altas: {high}<br/>
        • Médias: {stats.get('by_severity', {}).get('MEDIUM', 0)}<br/>
        • Baixas: {stats.get('by_severity', {}).get('LOW', 0)}
        """
        story.append(Paragraph(summary, styles['CustomBody']))
        story.append(Spacer(1, 15))
        
        # Estatísticas
        story.append(Paragraph("2. Estatísticas Detalhadas", styles['CustomHeading']))
        
        # Tabela de tipos de ataque
        if stats.get('by_type'):
            story.append(Paragraph("<b>Anomalias por Tipo:</b>", styles['CustomBody']))
            type_data = [['Tipo de Ataque', 'Quantidade', 'Percentagem']]
            for attack_type, count in sorted(stats['by_type'].items(), key=lambda x: -x[1]):
                pct = (count / total * 100) if total > 0 else 0
                type_data.append([attack_type, str(count), f"{pct:.1f}%"])
            
            type_table = Table(type_data, colWidths=[7*cm, 3*cm, 3*cm])
            type_table.setStyle(TableStyle([
                ('BACKGROUND', (0, 0), (-1, 0), colors.HexColor('#1e3a5f')),
                ('TEXTCOLOR', (0, 0), (-1, 0), colors.white),
                ('ALIGN', (0, 0), (-1, -1), 'CENTER'),
                ('FONTSIZE', (0, 0), (-1, -1), 9),
                ('GRID', (0, 0), (-1, -1), 0.5, colors.gray),
                ('ROWBACKGROUNDS', (0, 1), (-1, -1), [colors.white, colors.HexColor('#f3f4f6')]),
            ]))
            story.append(type_table)
            story.append(Spacer(1, 15))
        
        # Top IPs
        if stats.get('top_ips'):
            story.append(Paragraph("<b>Top IPs Suspeitos:</b>", styles['CustomBody']))
            ip_data = [['IP', 'Ocorrências']]
            for ip, count in stats['top_ips'][:10]:
                ip_data.append([ip, str(count)])
            
            ip_table = Table(ip_data, colWidths=[7*cm, 3*cm])
            ip_table.setStyle(TableStyle([
                ('BACKGROUND', (0, 0), (-1, 0), colors.HexColor('#1e3a5f')),
                ('TEXTCOLOR', (0, 0), (-1, 0), colors.white),
                ('ALIGN', (0, 0), (-1, -1), 'CENTER'),
                ('FONTSIZE', (0, 0), (-1, -1), 9),
                ('GRID', (0, 0), (-1, -1), 0.5, colors.gray),
            ]))
            story.append(ip_table)
            story.append(Spacer(1, 15))
        
        # Lista de Anomalias
        story.append(PageBreak())
        story.append(Paragraph("3. Anomalias Detetadas", styles['CustomHeading']))
        
        severity_colors = {
            'CRITICAL': colors.HexColor('#ef4444'),
            'HIGH': colors.HexColor('#f97316'),
            'MEDIUM': colors.HexColor('#eab308'),
            'LOW': colors.HexColor('#22c55e'),
        }
        
        for i, anomaly in enumerate(anomalies[:100], 1):  # Limitar a 100
            severity = anomaly.get('severity', 'MEDIUM')
            color = severity_colors.get(severity, colors.gray)
            
            anomaly_text = f"""
            <b>#{i} - {anomaly.get('type', 'UNKNOWN')}</b> 
            <font color="{color.hexval()}">[{severity}]</font><br/>
            <b>IP:</b> {anomaly.get('source_ip', 'N/A')} | 
            <b>Alvo:</b> {anomaly.get('target', 'N/A')}<br/>
            <b>Detalhe:</b> {anomaly.get('detail', 'N/A')[:200]}
            """
            story.append(Paragraph(anomaly_text, styles['CustomBody']))
            story.append(Spacer(1, 5))
        
        if len(anomalies) > 100:
            story.append(Paragraph(
                f"<i>... e mais {len(anomalies) - 100} anomalias (ver exportação completa em CSV/JSON)</i>",
                styles['CustomBody']
            ))
        
        # Recomendações
        story.append(PageBreak())
        story.append(Paragraph("4. Recomendações de Segurança", styles['CustomHeading']))
        
        recommendations = self._generate_recommendations(stats)
        for i, rec in enumerate(recommendations, 1):
            story.append(Paragraph(f"{i}. {rec}", styles['CustomBody']))
        
        # Footer
        story.append(Spacer(1, 30))
        story.append(HRFlowable(width="100%", thickness=1, color=colors.HexColor('#3b82f6')))
        story.append(Paragraph(
            f"<i>Gerado por {self.app_name} v{self.version} | {self.institution} | {datetime.now().year}</i>",
            ParagraphStyle('Footer', parent=styles['Normal'], fontSize=8, textColor=colors.gray, alignment=TA_CENTER)
        ))
        
        # Build PDF
        doc.build(story)
        return True
    
    # === CSV Export ===
    
    def _export_csv(self, anomalies: List[Dict], stats: Dict, filepath: Path) -> bool:
        """Exporta dados em CSV."""
        with open(filepath, 'w', newline='', encoding='utf-8') as f:
            if not anomalies:
                f.write("Sem anomalias detetadas\n")
                return True
            
            # Headers
            fieldnames = ['id', 'type', 'severity', 'source_ip', 'target', 
                         'detail', 'timestamp', 'score', 'log_file']
            
            writer = csv.DictWriter(f, fieldnames=fieldnames, extrasaction='ignore')
            writer.writeheader()
            
            for i, anomaly in enumerate(anomalies, 1):
                row = {'id': i}
                row.update(anomaly)
                writer.writerow(row)
        
        return True
    
    # === JSON Export ===
    
    def _export_json(self, anomalies: List[Dict], stats: Dict, filepath: Path) -> bool:
        """Exporta dados em JSON."""
        report = {
            'metadata': {
                'generated_at': datetime.now().isoformat(),
                'generator': f"{self.app_name} v{self.version}",
                'author': f"{self.author} (Nº {self.student_id})",
                'institution': self.institution,
            },
            'summary': {
                'total_entries_processed': stats.get('entries_processed', 0),
                'total_anomalies': stats.get('anomalies_detected', stats.get('total_anomalies', 0)),
                'by_severity': stats.get('by_severity', {}),
                'by_type': stats.get('by_type', {}),
                'top_ips': stats.get('top_ips', []),
            },
            'anomalies': anomalies,
            'recommendations': self._generate_recommendations(stats),
        }
        
        with open(filepath, 'w', encoding='utf-8') as f:
            json.dump(report, f, indent=2, ensure_ascii=False, default=str)
        
        return True
    
    # === DOCX Export ===
    
    def _export_docx(self, anomalies: List[Dict], stats: Dict, 
                     filepath: Path, title: str) -> bool:
        """Exporta relatório em DOCX."""
        if not DOCX_AVAILABLE:
            raise ImportError("python-docx não disponível")
        
        doc = Document()
        
        # Title
        title_para = doc.add_heading("🦉 LOG SENTINEL", 0)
        title_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
        
        subtitle = doc.add_heading(title, level=1)
        subtitle.alignment = WD_ALIGN_PARAGRAPH.CENTER
        
        # Meta info
        doc.add_paragraph(f"Data: {datetime.now().strftime('%d/%m/%Y %H:%M')}")
        doc.add_paragraph(f"Autor: {self.author} (Nº {self.student_id})")
        doc.add_paragraph(f"Instituição: {self.institution}")
        doc.add_paragraph("")
        
        # Summary
        doc.add_heading("1. Resumo Executivo", level=2)
        total = stats.get('total_anomalies', stats.get('anomalies_detected', 0))
        doc.add_paragraph(
            f"Esta análise processou {stats.get('entries_processed', 'N/A')} entradas "
            f"e identificou {total} anomalias de segurança."
        )
        
        # Stats table
        doc.add_heading("2. Estatísticas", level=2)
        
        if stats.get('by_severity'):
            doc.add_paragraph("Anomalias por Severidade:")
            table = doc.add_table(rows=1, cols=2)
            table.style = 'Table Grid'
            hdr_cells = table.rows[0].cells
            hdr_cells[0].text = 'Severidade'
            hdr_cells[1].text = 'Quantidade'
            
            for severity, count in stats['by_severity'].items():
                row_cells = table.add_row().cells
                row_cells[0].text = severity
                row_cells[1].text = str(count)
        
        # Anomalies
        doc.add_heading("3. Anomalias Detetadas", level=2)
        
        for i, anomaly in enumerate(anomalies[:50], 1):
            p = doc.add_paragraph()
            p.add_run(f"#{i} - {anomaly.get('type', 'UNKNOWN')} ").bold = True
            p.add_run(f"[{anomaly.get('severity', 'MEDIUM')}]")
            
            doc.add_paragraph(
                f"IP: {anomaly.get('source_ip', 'N/A')} | "
                f"Detalhe: {anomaly.get('detail', 'N/A')[:150]}"
            )
        
        # Recommendations
        doc.add_heading("4. Recomendações", level=2)
        for rec in self._generate_recommendations(stats):
            doc.add_paragraph(rec, style='List Bullet')
        
        # Save
        doc.save(str(filepath))
        return True
    
    def _generate_recommendations(self, stats: Dict) -> List[str]:
        """Gera recomendações baseadas nos resultados."""
        recommendations = []
        
        by_type = stats.get('by_type', {})
        by_severity = stats.get('by_severity', {})
        
        # Recomendações baseadas em tipos de ataque
        if by_type.get('SQL_INJECTION', 0) > 0:
            recommendations.append(
                "Implementar prepared statements e validação de input em todas as queries SQL."
            )
        
        if by_type.get('XSS', 0) > 0:
            recommendations.append(
                "Implementar sanitização de output e Content Security Policy (CSP)."
            )
        
        if by_type.get('BRUTE_FORCE', 0) > 0:
            recommendations.append(
                "Implementar rate limiting, CAPTCHA e autenticação multi-fator."
            )
        
        if by_type.get('PATH_TRAVERSAL', 0) > 0:
            recommendations.append(
                "Validar e sanitizar todos os caminhos de ficheiros no servidor."
            )
        
        if by_type.get('COMMAND_INJECTION', 0) > 0:
            recommendations.append(
                "Evitar execução de comandos do sistema com input de utilizador."
            )
        
        if by_type.get('SCANNER', 0) > 0:
            recommendations.append(
                "Implementar Web Application Firewall (WAF) para bloquear scanners."
            )
        
        if by_type.get('DDOS', 0) > 0:
            recommendations.append(
                "Configurar proteção DDoS e rate limiting no servidor/CDN."
            )
        
        # Recomendações baseadas em severidade
        if by_severity.get('CRITICAL', 0) > 0:
            recommendations.insert(0, 
                "⚠️ URGENTE: Investigar imediatamente as anomalias críticas detetadas."
            )
        
        # Recomendações gerais
        if stats.get('top_ips'):
            recommendations.append(
                f"Considerar bloquear os IPs mais ativos: {', '.join([ip for ip, _ in stats['top_ips'][:3]])}"
            )
        
        recommendations.append("Manter todos os sistemas e dependências atualizados.")
        recommendations.append("Implementar logging centralizado e monitorização contínua.")
        recommendations.append("Realizar testes de penetração periódicos.")
        
        return recommendations


# Teste do módulo
if __name__ == "__main__":
    print("🔧 Teste do ReportExporter")
    
    exporter = ReportExporter()
    
    # Dados de teste
    test_anomalies = [
        {'type': 'SQL_INJECTION', 'severity': 'CRITICAL', 'source_ip': '192.168.1.100', 
         'detail': 'Tentativa de SQL injection detetada', 'timestamp': datetime.now().isoformat()},
        {'type': 'XSS', 'severity': 'HIGH', 'source_ip': '10.0.0.50',
         'detail': 'Script injection em parâmetro', 'timestamp': datetime.now().isoformat()},
    ]
    
    test_stats = {
        'entries_processed': 10000,
        'anomalies_detected': 2,
        'by_type': {'SQL_INJECTION': 1, 'XSS': 1},
        'by_severity': {'CRITICAL': 1, 'HIGH': 1},
        'top_ips': [('192.168.1.100', 5), ('10.0.0.50', 3)],
    }
    
    # Testar JSON (sempre disponível)
    exporter.export(test_anomalies, test_stats, "test_report.json", "json")
    print("✅ JSON exportado")
    
    # Testar CSV
    exporter.export(test_anomalies, test_stats, "test_report.csv", "csv")
    print("✅ CSV exportado")
    
    # Limpar
    import os
    for f in ["test_report.json", "test_report.csv"]:
        if os.path.exists(f):
            os.remove(f)
    
    print("✅ Teste concluído!")
